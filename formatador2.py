import customtkinter as ctk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx2pdf import convert
import os
from datetime import datetime
import platform


def ler_conteudo(caminho):
    if caminho.endswith(".txt"):
        with open(caminho, "r", encoding="utf-8") as f:
            return f.read()
    elif caminho.endswith(".docx"):
        doc = Document(caminho)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        raise ValueError("Formato de arquivo não suportado.")


def adicionar_capa(doc, titulo, autor):
    doc.add_paragraph("")
    p1 = doc.add_paragraph(autor)
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    p2 = doc.add_paragraph(titulo.upper())
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.runs[0].bold = True

    doc.add_paragraph("")
    data = datetime.now().strftime("%d de %B de %Y")
    p3 = doc.add_paragraph(data)
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()


def adicionar_sumario(doc):
    p = doc.add_paragraph("Sumário", style='Heading 1')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(">>> O sumário será gerado automaticamente ao abrir no Word. <<<")
    doc.add_page_break()


def formatar_abnt(texto, nome_base, titulo, autor):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    adicionar_capa(doc, titulo=titulo, autor=autor)
    adicionar_sumario(doc)

    for linha in texto.split('\n'):
        if linha.strip():
            if linha.strip().lower().startswith("capítulo"):
                doc.add_paragraph(linha.strip(), style='Heading 1')
            else:
                p = doc.add_paragraph()
                p.paragraph_format.first_line_indent = Cm(1.25)
                p.paragraph_format.line_spacing = 1.5
                p.alignment = 3
                p.add_run(linha.strip())

    nome_docx = f"{nome_base}_formatado.docx"
    nome_pdf = f"{nome_base}_formatado.pdf"
    doc.save(nome_docx)

    if platform.system() == "Windows":
        try:
            convert(nome_docx, nome_pdf)
            return True, nome_pdf
        except Exception as e:
            mensagem = str(e)
            if "Word" in mensagem or "Office" in mensagem:
                mensagem += "\n⚠️ Word é necessário no Windows para converter para PDF."
            return False, f"Erro ao converter PDF:\n{mensagem}"
    else:
        return False, "⚠️ PDF só pode ser gerado automaticamente no Windows com Word instalado."


def selecionar_arquivo():
    caminho = filedialog.askopenfilename(filetypes=[("Texto ou Word", "*.txt *.docx")])
    if caminho:
        entrada_arquivo.delete(0, "end")
        entrada_arquivo.insert(0, caminho)


def gerar_documento():
    caminho = entrada_arquivo.get()
    titulo = entrada_titulo.get().strip()
    autor = entrada_autor.get().strip()

    if not caminho or not os.path.exists(caminho):
        messagebox.showerror("Erro", "Selecione um arquivo válido.")
        return
    if not titulo or not autor:
        messagebox.showerror("Erro", "Preencha o título e o nome do autor.")
        return

    try:
        progresso.set(0.2)
        status_log.configure(text="📂 Lendo conteúdo do arquivo...")
        app.update()

        texto = ler_conteudo(caminho)
        nome_base = os.path.splitext(os.path.basename(caminho))[0]

        progresso.set(0.5)
        status_log.configure(text="📄 Formatando documento ABNT...")
        app.update()

        sucesso, saida = formatar_abnt(texto, nome_base, titulo, autor)

        progresso.set(1.0)
        if sucesso:
            status_log.configure(text="✅ Documento gerado com sucesso!")
            messagebox.showinfo("Sucesso", f"PDF gerado:\n{saida}")
        else:
            status_log.configure(text="⚠️ PDF não gerado, mas DOCX salvo.")
            messagebox.showwarning("Aviso", f"Documento Word salvo.\n{saida}")
    except Exception as e:
        progresso.set(0)
        status_log.configure(text="❌ Ocorreu um erro.")
        messagebox.showerror("Erro", str(e))


# === Interface com CustomTkinter ===
ctk.set_appearance_mode("dark")  # Pode trocar por "dark"
ctk.set_default_color_theme("dark-blue")  # Temas: "blue", "dark-blue", "green"

app = ctk.CTk()
app.title("Formatador ABNT Moderno")
app.geometry("550x540")
app.resizable(False, False)

frame = ctk.CTkFrame(app, corner_radius=20)
frame.pack(padx=30, pady=20, fill="both", expand=True)

titulo_label = ctk.CTkLabel(frame, text="Formatador ABNT", font=ctk.CTkFont(size=22, weight="bold"))
titulo_label.pack(pady=10)

entrada_arquivo = ctk.CTkEntry(frame, placeholder_text="Caminho do arquivo", width=400, height=40, corner_radius=15)
entrada_arquivo.pack(pady=10)

botao_arquivo = ctk.CTkButton(frame, text="Selecionar Arquivo", command=selecionar_arquivo, corner_radius=15)
botao_arquivo.pack(pady=5)

entrada_titulo = ctk.CTkEntry(frame, placeholder_text="Título do Trabalho", width=400, height=40, corner_radius=15)
entrada_titulo.pack(pady=10)

entrada_autor = ctk.CTkEntry(frame, placeholder_text="Nome do Autor", width=400, height=40, corner_radius=15)
entrada_autor.pack(pady=10)

botao_gerar = ctk.CTkButton(frame, text="Gerar Documento", command=gerar_documento, width=200, height=45, corner_radius=15)
botao_gerar.pack(pady=10)

progresso = ctk.CTkProgressBar(frame, width=300)
progresso.set(0)
progresso.pack(pady=(10, 5))

status_log = ctk.CTkLabel(frame, text="", text_color="gray")
status_log.pack(pady=5)

# Modo claro/escuro
def alternar_modo():
    modo_atual = ctk.get_appearance_mode()
    novo_modo = "dark" if modo_atual == "light" else "light"
    ctk.set_appearance_mode(novo_modo)
    botao_modo.configure(text="🌙 Modo Claro" if novo_modo == "dark" else "🌞 Modo Escuro")

botao_modo = ctk.CTkButton(frame, text="🌞 Modo Escuro", command=alternar_modo, width=150, corner_radius=15)
botao_modo.pack(pady=10)

app.mainloop()
