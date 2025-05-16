import tkinter as tk
from tkinter import filedialog, messagebox
from docx import Document
from docx.shared import Cm, Pt
from docx2pdf import convert
import os

def formatar_abnt(texto, nome_docx='documento_abnt.docx', nome_pdf='documento_abnt.pdf'):
    doc = Document()

    # Configurar margens
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(3)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)

    # Estilo da fonte
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    for linha in texto.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.25)
        p.paragraph_format.line_spacing = 1.5
        p.alignment = 3  # Justificado
        run = p.add_run(linha)

    # Salvar .docx
    doc.save(nome_docx)

    # Tentar converter para PDF
    try:
        convert(nome_docx, nome_pdf)
        return True, nome_pdf
    except Exception as e:
        return False, str(e)

def selecionar_arquivo():
    caminho = filedialog.askopenfilename(filetypes=[("Arquivos de texto", "*.txt")])
    if caminho:
        try:
            with open(caminho, "r", encoding="utf-8") as f:
                conteudo = f.read()
            nome_base = os.path.splitext(os.path.basename(caminho))[0]
            docx_nome = f"{nome_base}_formatado.docx"
            pdf_nome = f"{nome_base}_formatado.pdf"

            sucesso, resultado = formatar_abnt(conteudo, docx_nome, pdf_nome)

            if sucesso:
                messagebox.showinfo("Sucesso", f"PDF gerado: {resultado}")
            else:
                messagebox.showwarning("Erro ao gerar PDF", f"Docx gerado: {docx_nome}\nErro no PDF: {resultado}")
        except Exception as e:
            messagebox.showerror("Erro", f"Erro ao processar o arquivo:\n{str(e)}")

# Interface com Tkinter
janela = tk.Tk()
janela.title("Formatador ABNT - PDF Automático")
janela.geometry("400x200")

label = tk.Label(janela, text="Selecione um arquivo .txt para formatar segundo a ABNT", wraplength=380)
label.pack(pady=20)

botao = tk.Button(janela, text="Selecionar Arquivo", command=selecionar_arquivo)
botao.pack(pady=10)

janela.mainloop()
